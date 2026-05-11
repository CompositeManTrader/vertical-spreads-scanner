"""
Utilidades reusables para construir reportes Word con python-docx.
"""

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def new_document(title: str, subtitle: str = "") -> Document:
    doc = Document()
    # Margenes
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Titulo
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(11)

    # Fecha
    p = doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()
    return doc


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_callout(doc: Document, text: str, color: str = "warning") -> None:
    """Bloque destacado para advertencias / hallazgos importantes."""
    colors = {
        "warning": RGBColor(0xCC, 0x66, 0x00),
        "info": RGBColor(0x00, 0x66, 0xCC),
        "success": RGBColor(0x00, 0x88, 0x44),
        "danger": RGBColor(0xCC, 0x00, 0x00),
    }
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = colors.get(color, RGBColor(0, 0, 0))


def add_table_from_df(
    doc: Document,
    df: pd.DataFrame,
    style: str = "Light Grid Accent 1",
    float_format: str = "{:.4f}",
    int_format: str = "{:,}",
) -> None:
    """Inserta una tabla con un DataFrame, formateando floats y enteros."""
    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    try:
        table.style = style
    except KeyError:
        table.style = "Light Grid"

    # Header
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = str(c)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    # Body
    for ri, (_, row) in enumerate(df.iterrows(), start=1):
        for ci, c in enumerate(cols):
            v = row[c]
            if isinstance(v, float):
                txt = float_format.format(v) if pd.notna(v) else ""
            elif isinstance(v, (int,)) and not isinstance(v, bool):
                txt = int_format.format(v)
            elif pd.isna(v):
                txt = ""
            else:
                txt = str(v)
            cell = table.rows[ri].cells[ci]
            cell.text = txt
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def add_image(doc: Document, image_path: str | Path, width_inches: float = 6.0) -> None:
    doc.add_picture(str(image_path), width=Inches(width_inches))


def save(doc: Document, path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
